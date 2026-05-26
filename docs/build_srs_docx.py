from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path("docs/SRS_PotionCheck_IEEE.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    run.font.name = "Times New Roman"
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def style_doc(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)

    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def setup_section(section, title="Software Requirements Specification for PotionCheck"):
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
    add_page_number(section.footer.paragraphs[0])


def para(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "D9EAF7")
        set_cell_text(hdr[i], h, True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return t


def heading(doc, level, text):
    doc.add_heading(text, level=level)


def feature(doc, number, title, description, priority, sequence, requirements):
    heading(doc, 2, f"4.{number} {title}")
    heading(doc, 3, f"4.{number}.1 Description and Priority")
    para(doc, f"{description} Priority: {priority}.")
    heading(doc, 3, f"4.{number}.2 Stimulus/Response Sequences")
    table(doc, ["Step", "User Action or System Event", "System Response"], sequence)
    heading(doc, 3, f"4.{number}.3 Functional Requirements")
    table(doc, ["ID", "Requirement"], requirements)


def build():
    doc = Document()
    style_doc(doc)
    setup_section(doc.sections[0], "")

    para(doc, "Copyright (c) 1999 by Karl E. Wiegers. Permission is granted to use, modify, and distribute this document.", WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()
    doc.add_paragraph()
    p = para(doc, "Software Requirements\nSpecification\nfor\nPotionCheck", WD_ALIGN_PARAGRAPH.CENTER, True)
    for run in p.runs:
        run.font.size = Pt(24)
    doc.add_paragraph()
    para(doc, "Version 1.0 approved", WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Prepared by Nishant (Roll No. 2310993892) and Mohit (Roll No. 2310993967)", WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "PBL 6 Project Team", WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, date.today().strftime("%d %B %Y"), WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    setup_section(doc.sections[-1])
    heading(doc, 1, "Table of Contents")
    toc_rows = [
        ("Revision History", "ii"),
        ("1. Introduction", "1"),
        ("1.1 Purpose", "1"),
        ("1.2 Document Conventions", "1"),
        ("1.3 Intended Audience and Reading Suggestions", "1"),
        ("1.4 Product Scope", "2"),
        ("1.5 References", "2"),
        ("2. Overall Description", "3"),
        ("2.1 Product Perspective", "3"),
        ("2.2 Product Functions", "3"),
        ("2.3 User Classes and Characteristics", "4"),
        ("2.4 Operating Environment", "4"),
        ("2.5 Design and Implementation Constraints", "5"),
        ("2.6 User Documentation", "5"),
        ("2.7 Assumptions and Dependencies", "6"),
        ("3. External Interface Requirements", "6"),
        ("3.1 User Interfaces", "6"),
        ("3.2 Hardware Interfaces", "7"),
        ("3.3 Software Interfaces", "7"),
        ("3.4 Communications Interfaces", "8"),
        ("4. System Features", "9"),
        ("5. Other Nonfunctional Requirements", "16"),
        ("6. Other Requirements", "19"),
        ("Appendix A: Glossary", "21"),
        ("Appendix B: Analysis Models", "22"),
        ("Appendix C: To Be Determined List", "23"),
    ]
    for title, page in toc_rows:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.3))
        run = p.add_run(f"{title}\t{page}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    heading(doc, 1, "Revision History")
    table(
        doc,
        ["Name", "Date", "Reason For Changes", "Version"],
        [["Nishant and Mohit", date.today().strftime("%d %B %Y"), "Initial SRS prepared for PotionCheck using the IEEE-style template.", "1.0"]],
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    setup_section(doc.sections[-1])
    heading(doc, 1, "1. Introduction")
    heading(doc, 2, "1.1 Purpose")
    para(doc, "This Software Requirements Specification describes the requirements for PotionCheck, version 1.0. PotionCheck is an AI-powered ingredient intelligence scanner made for the PBL 6 project. It helps users understand packaged food products by scanning barcodes, uploading food label images, pasting ingredient text, or selecting simple food items for analysis.")
    para(doc, "The document covers the complete software system: the React/Vite frontend, the FastAPI backend, authentication, user profile management, barcode lookup, OCR processing, AI ingredient analysis, scan history, chatbot support, health checks, deployment configuration, and data storage. It is written in clear language so developers, testers, faculty members, and future maintainers can understand the project without guessing the system behavior.")

    heading(doc, 2, "1.2 Document Conventions")
    para(doc, "This document follows the IEEE-style SRS structure provided in the template. Section numbers are used for traceability. Functional requirements use the tag FR followed by a three-digit number. Nonfunctional requirements use the tag NFR followed by a three-digit number.")
    bullets(doc, [
        "The word shall means the requirement is mandatory.",
        "The word should means the requirement is recommended when practical.",
        "API paths, file paths, commands, and configuration keys are written in plain technical form.",
        "Each system feature includes description and priority, stimulus/response sequence, and functional requirements.",
        "Body text is justified to match the formal written style of the source template.",
    ])

    heading(doc, 2, "1.3 Intended Audience and Reading Suggestions")
    para(doc, "This document is intended for project evaluators, faculty members, developers, testers, deployment maintainers, and future contributors. Evaluators can use it to check whether the project scope is complete. Developers can use it to understand what must be implemented. Testers can use it to prepare test cases. Maintainers can use it to understand the architecture and deployment needs.")
    numbered(doc, [
        "Read Section 1 for purpose, scope, conventions, audience, and references.",
        "Read Section 2 for the overall view of the product, users, environment, constraints, assumptions, and dependencies.",
        "Read Section 3 for interface requirements between the user, hardware, software services, and communication protocols.",
        "Read Section 4 for detailed system features and functional requirements.",
        "Read Section 5 for performance, safety, security, quality, and business rules.",
        "Read Section 6 and the appendices for database, deployment, glossary, analysis model, and remaining TBD items.",
    ])

    heading(doc, 2, "1.4 Product Scope")
    para(doc, "PotionCheck converts food product information into simple, personalized health guidance. The system combines barcode product lookup, OCR, AI reasoning, user health profiles, and scan history in one web application. The main objective is to help users make better food choices by explaining ingredients, nutrition values, possible risks, and suitability for their profile.")
    bullets(doc, [
        "Analyze products from barcode, uploaded image, pasted text, and selected produce.",
        "Detect caution-worthy ingredients and explain them in easy language.",
        "Personalize warnings using allergies, health conditions, and diet type.",
        "Show safety score, verdict, health score, nutrition observations, AI summary, recommendation, frequency advice, and ingredient-level explanations.",
        "Save scan history for authenticated users and recent local history for quick access.",
        "Use Groq as the primary AI provider, Ollama as fallback, and local rule-based analysis when AI services are unavailable.",
    ])
    para(doc, "PotionCheck is not a medical diagnosis system. It is an educational and informational project. Users with serious allergies, pregnancy, chronic diseases, or strict diets should verify product labels and consult a qualified professional when needed.")

    heading(doc, 2, "1.5 References")
    table(doc, ["Reference", "Location or Source"], [
        ["IEEE-style SRS template", r"C:\Users\LOQ\Downloads\srs_template-ieee.pdf"],
        ["Project deployment guide", "DEPLOYMENT.md"],
        ["Docker Compose configuration", "docker-compose.yml"],
        ["Render deployment configuration", "render.yaml"],
        ["Frontend package definition", "frontend/package.json"],
        ["Backend dependencies", "backend/requirements.txt"],
        ["Frontend routes", "frontend/src/App.jsx"],
        ["Backend entry point", "backend/app/main.py"],
        ["Backend routers", "backend/app/api/"],
        ["Database models", "backend/app/models/entities.py"],
        ["AI analysis service", "backend/app/services/ai.py"],
        ["Open Food Facts service", "backend/app/services/open_food_facts.py"],
        ["Open Food Facts API", "https://world.openfoodfacts.org"],
        ["Groq API", "Groq chat completion service"],
        ["Ollama", "Local model runtime used as fallback"],
    ])

    heading(doc, 1, "2. Overall Description")
    heading(doc, 2, "2.1 Product Perspective")
    para(doc, "PotionCheck is a new, self-contained web application. It is not a replacement for an existing system. The frontend is a React single-page application and the backend is a FastAPI REST API. The backend connects to database storage, cache storage, Open Food Facts, Groq, Ollama, OCR services, and local rule-based analysis.")
    para(doc, "The system is divided into clear parts: user interface, API layer, authentication layer, scanner and OCR layer, AI analysis layer, history layer, and deployment layer. The user interacts only with the web interface, while the frontend communicates with the backend through JSON and multipart HTTP requests.")

    heading(doc, 2, "2.2 Product Functions")
    bullets(doc, [
        "Register, login, refresh session, logout, and retrieve current user.",
        "Create and update health profile with allergies, health conditions, and diet type.",
        "Scan product barcode using camera, upload barcode image, or enter barcode manually.",
        "Fetch product information from Open Food Facts.",
        "Upload a food label image and extract ingredient text through OCR.",
        "Clean and parse pasted ingredient text.",
        "Analyze ingredients and nutrition using AI with fallback logic.",
        "Display report with product details, ingredient cards, verdict, score, warnings, summary, and recommendation.",
        "Save, search, filter, and delete scan history.",
        "Provide chatbot and health-monitoring support.",
    ])

    heading(doc, 2, "2.3 User Classes and Characteristics")
    table(doc, ["User Class", "Description", "Expected Skill Level"], [
        ["Guest User", "Uses scanning and analysis without signing in. Local profile and recent history may be stored in the browser.", "Basic web usage"],
        ["Registered User", "Creates an account, saves profile, receives personalized analysis, and accesses server-side history.", "Basic web usage"],
        ["Health-Conscious Consumer", "Reviews packaged food ingredients and nutrition before eating or buying.", "Basic nutrition awareness"],
        ["Developer or Maintainer", "Runs, debugs, deploys, and extends the application.", "Web development knowledge"],
        ["Evaluator or Instructor", "Reviews project scope, implementation, and documentation.", "General technical review"],
    ])

    heading(doc, 2, "2.4 Operating Environment")
    bullets(doc, [
        "Frontend runs in modern browsers such as Chrome or Edge with JavaScript enabled.",
        "Camera scanning requires localhost or HTTPS and browser camera permission.",
        "Frontend development uses Node.js, Vite, React, and Tailwind CSS.",
        "Backend runs on Python with FastAPI and Uvicorn.",
        "Database uses SQLite through SQLAlchemy async models.",
        "Redis is used for product lookup cache when available.",
        "Open Food Facts network access is required for barcode product lookup.",
        "Groq API access is optional but preferred for AI analysis.",
        "Ollama local server is optional fallback for AI analysis.",
        "Deployment can run locally, through Docker Compose, or through Render.",
    ])

    heading(doc, 2, "2.5 Design and Implementation Constraints")
    bullets(doc, [
        "The frontend shall remain a React single-page application using React Router.",
        "The backend shall expose REST APIs under the /api path.",
        "Authentication shall use JWT access tokens and HTTP-only refresh cookies.",
        "Passwords shall be stored only after bcrypt hashing.",
        "Barcode lookup depends on Open Food Facts coverage and availability.",
        "Groq analysis requires a valid GROQ_API_KEY environment variable.",
        "Ollama fallback requires a reachable Ollama server.",
        "Uploaded label images shall be limited to image files and a maximum size of 10 MB.",
        "SQLite is acceptable for project scale but may require migration for heavy production traffic.",
    ])

    heading(doc, 2, "2.6 User Documentation")
    para(doc, "User documentation is provided through interface labels, toast notifications, validation messages, scan status messages, and project deployment files. Additional documentation should explain how to run the app, configure environment variables, scan a barcode, upload an ingredient label, paste ingredient text, and understand scores and warnings.")

    heading(doc, 2, "2.7 Assumptions and Dependencies")
    table(doc, ["Type", "Item"], [
        ["Assumption", "Users understand that PotionCheck gives informational guidance and not medical advice."],
        ["Assumption", "Product data from Open Food Facts can be incomplete or inaccurate."],
        ["Assumption", "AI output may need human judgment and should not be treated as clinical diagnosis."],
        ["Assumption", "Users grant camera permission when using live barcode scanning."],
        ["Dependency", "Open Food Facts provides barcode product information."],
        ["Dependency", "Groq provides primary LLM analysis when configured."],
        ["Dependency", "Ollama provides fallback LLM analysis when reachable."],
        ["Dependency", "Tesseract OCR extracts text from label images."],
        ["Dependency", "Redis supports cached product lookup data."],
        ["Dependency", "SQLite stores users, profiles, scan history, and analysis data."],
    ])

    heading(doc, 1, "3. External Interface Requirements")
    heading(doc, 2, "3.1 User Interfaces")
    table(doc, ["Page", "Route", "Purpose"], [
        ["Landing", "/", "Introduce the app and route users to major features."],
        ["Profile", "/profile", "Register, login, and manage health profile."],
        ["Scanner", "/scanner", "Scan barcode, upload label, paste text, or select produce."],
        ["Analysis", "/analysis/:productId", "Display generated ingredient and nutrition report."],
        ["History", "/history", "View saved scan history."],
        ["Chatbot", "/chatbot", "Ask food and ingredient questions."],
        ["About", "/about", "Explain application purpose."],
        ["Developer", "/developer", "Present developer and project information."],
    ])
    bullets(doc, [
        "The UI shall show success, warning, and error toast messages.",
        "The UI shall show progress indicators during OCR, lookup, and AI analysis.",
        "The UI shall be responsive for desktop and mobile screens.",
        "The UI shall display clear messages when camera permission, network, or backend access fails.",
        "The analysis report shall use easy language so a normal consumer can understand the result.",
    ])

    heading(doc, 2, "3.2 Hardware Interfaces")
    bullets(doc, [
        "Camera is required for live barcode scanning.",
        "Keyboard is required for manual barcode entry, login, profile entry, and pasted ingredient text.",
        "File input is required for uploading barcode images and food label images.",
        "Network interface is required for backend API calls, Open Food Facts, Groq, and deployment access.",
    ])

    heading(doc, 2, "3.3 Software Interfaces")
    table(doc, ["Software Component", "Purpose"], [
        ["React 18", "Frontend user interface."],
        ["React Router", "Client-side routing."],
        ["Vite", "Frontend development and build tooling."],
        ["Tailwind CSS", "Frontend styling."],
        ["Framer Motion and Three.js", "Animations and visual effects."],
        ["Zustand", "Frontend state management."],
        ["Axios", "HTTP client for API communication."],
        ["html5-qrcode", "Barcode scanning in the browser."],
        ["Tesseract.js and pytesseract", "OCR support for label text extraction."],
        ["FastAPI and Uvicorn", "Backend REST API server."],
        ["SQLAlchemy async and aiosqlite", "Database access."],
        ["Redis", "Cache storage."],
        ["Groq SDK", "Primary AI provider."],
        ["Ollama", "Local fallback AI provider."],
        ["Open Food Facts API", "Barcode product data source."],
    ])

    heading(doc, 2, "3.4 Communications Interfaces")
    para(doc, "The frontend communicates with the backend through HTTP or HTTPS REST API calls. JSON is used for normal request and response bodies. Multipart form data is used for image uploads. The backend communicates with Open Food Facts and Groq through HTTPS. Ollama communication uses the configured local or network base URL.")
    table(doc, ["Method", "Endpoint", "Purpose"], [
        ["POST", "/api/auth/register", "Register user."],
        ["POST", "/api/auth/login", "Login user."],
        ["POST", "/api/auth/refresh", "Refresh access token."],
        ["POST", "/api/auth/logout", "Logout user."],
        ["GET", "/api/auth/me", "Retrieve current user."],
        ["GET", "/api/profile/", "Get user profile."],
        ["POST", "/api/profile/", "Save user profile."],
        ["POST", "/api/scanner/barcode", "Fetch product by barcode."],
        ["POST", "/api/scanner/barcode/analyze", "Analyze barcode product."],
        ["POST", "/api/scanner/upload", "Upload image for OCR."],
        ["POST", "/api/scanner/text", "Clean and parse ingredient text."],
        ["POST", "/api/analysis/analyze", "Analyze custom ingredient payload."],
        ["GET", "/api/analysis/{analysis_id}", "Get saved analysis."],
        ["GET", "/api/analysis/product/{barcode}", "Get latest analysis for barcode."],
        ["GET", "/api/history/", "Get scan history."],
        ["DELETE", "/api/history/{scan_id}", "Delete one scan."],
        ["DELETE", "/api/history/all", "Delete all scans."],
        ["GET", "/api/health/", "Service health check."],
    ])

    heading(doc, 1, "4. System Features")
    para(doc, "This section lists the major services provided by PotionCheck. Each feature includes description and priority, stimulus/response sequence, and functional requirements.")

    feature(doc, 1, "Authentication and User Session Management",
            "The system lets users register, log in, refresh tokens, log out, and retrieve the current authenticated user.",
            "High",
            [["1", "User submits registration details.", "System validates email and password, creates account, hashes password, and returns token."],
             ["2", "User submits login details.", "System checks credentials and returns access token plus refresh cookie."],
             ["3", "Access token expires.", "System refreshes token when a valid refresh cookie is present."],
             ["4", "User logs out.", "System clears refresh cookie."]],
            [["FR-001", "The system shall allow a new user to register with email, password, and full name."],
             ["FR-002", "The system shall reject registration when the email is already registered."],
             ["FR-003", "The system shall validate email format and password length."],
             ["FR-004", "The system shall allow users to log in with valid credentials."],
             ["FR-005", "The system shall reject invalid login credentials."],
             ["FR-006", "The system shall issue JWT access tokens after successful login or registration."],
             ["FR-007", "The system shall store refresh tokens in HTTP-only cookies."],
             ["FR-008", "The system shall allow token refresh with valid refresh cookies."],
             ["FR-009", "The system shall allow users to log out and clear refresh cookies."],
             ["FR-010", "The system shall expose a current-user endpoint for authenticated users."]])

    feature(doc, 2, "Health Profile Management",
            "The system stores user allergies, health conditions, and diet type for personalized ingredient analysis.",
            "High",
            [["1", "User opens profile page.", "System loads existing profile if signed in or local profile if anonymous."],
             ["2", "User saves allergies, conditions, and diet type.", "System validates values and stores the profile."],
             ["3", "User analyzes product.", "System uses the profile to personalize warnings."]],
            [["FR-011", "The system shall allow authenticated users to retrieve their saved profile."],
             ["FR-012", "The system shall allow authenticated users to save allergies, health conditions, and diet type."],
             ["FR-013", "The system shall validate profile values against allowed lists."],
             ["FR-014", "The system shall support allergies such as nuts, gluten, dairy, soy, eggs, shellfish, peanuts, and fish."],
             ["FR-015", "The system shall support health conditions such as diabetes, hypertension, heart disease, pregnancy, kidney disease, celiac disease, and lactose intolerance."],
             ["FR-016", "The frontend shall persist a local profile for anonymous use."]])

    feature(doc, 3, "Barcode Scanning and Product Lookup",
            "The system analyzes products using live camera barcode scanning, uploaded barcode images, or manual barcode entry.",
            "High",
            [["1", "User scans or enters a barcode.", "System validates barcode format."],
             ["2", "Barcode is valid.", "Backend fetches product data from Open Food Facts or cache."],
             ["3", "Product has ingredients.", "System starts analysis."],
             ["4", "Product is missing or incomplete.", "System returns a clear error message."]],
            [["FR-017", "The system shall support manual barcode entry for 8 to 13 digit UPC/EAN codes."],
             ["FR-018", "The system shall validate barcode format before analysis requests."],
             ["FR-019", "The frontend shall support live barcode scanning using browser camera access."],
             ["FR-020", "The frontend shall support barcode detection from uploaded images."],
             ["FR-021", "The backend shall fetch product details from Open Food Facts by barcode."],
             ["FR-022", "The backend shall cache product lookup results."],
             ["FR-023", "The backend shall return 404 when a barcode product is not found."],
             ["FR-024", "The backend shall return a clear error when Open Food Facts lookup fails."],
             ["FR-025", "The system shall analyze barcode products when ingredients are available."],
             ["FR-026", "The system shall notify the user if a product exists but ingredient data is missing."]])

    feature(doc, 4, "Label Upload and OCR",
            "The system allows users to upload food label images and extract ingredient text for analysis.",
            "High",
            [["1", "User uploads a label image.", "System checks file type and size."],
             ["2", "Image is accepted.", "Backend stores the file temporarily and runs OCR."],
             ["3", "Text is extracted.", "Frontend displays it for user review and analysis."],
             ["4", "OCR fails.", "System returns an empty or clear result without crashing."]],
            [["FR-027", "The frontend shall allow image upload for product label OCR."],
             ["FR-028", "The backend shall accept image uploads only."],
             ["FR-029", "The backend shall reject files larger than 10 MB."],
             ["FR-030", "The backend shall store uploaded images temporarily in the uploads directory."],
             ["FR-031", "The backend shall run OCR on uploaded images."],
             ["FR-032", "The frontend shall display extracted text for user review and editing."],
             ["FR-033", "The system shall allow analysis of OCR-extracted ingredients."]])

    feature(doc, 5, "Pasted Ingredient Text Analysis",
            "The system lets users paste ingredient text directly when barcode or image input is not available.",
            "Medium",
            [["1", "User pastes ingredient text.", "System cleans and parses the text."],
             ["2", "User requests analysis.", "Backend validates text and sends it for analysis."],
             ["3", "Text is empty.", "System asks the user to provide valid ingredient text."]],
            [["FR-034", "The frontend shall provide a text area for pasted ingredient lists."],
             ["FR-035", "The frontend shall support basic ingredient parsing for preview."],
             ["FR-036", "The backend shall validate that ingredient text is not empty."],
             ["FR-037", "The backend shall analyze pasted ingredient text without requiring barcode data."]])

    feature(doc, 6, "AI Ingredient and Nutrition Analysis",
            "The system generates ingredient and nutrition analysis using Groq first, Ollama as fallback, and local rules as the final fallback.",
            "High",
            [["1", "System receives product/profile data.", "Backend prepares an analysis request."],
             ["2", "Groq is configured and available.", "Backend uses Groq for AI analysis."],
             ["3", "Groq fails or is not configured.", "Backend tries Ollama."],
             ["4", "Ollama is unavailable.", "Backend produces local rule-based analysis."],
             ["5", "AI output has missing fields.", "Backend normalizes the response."]],
            [["FR-038", "The backend shall analyze ingredients using the configured AI provider."],
             ["FR-039", "The default AI provider shall be Groq when configured."],
             ["FR-040", "The backend shall fall back to Ollama when Groq is unavailable, rate-limited, or not configured."],
             ["FR-041", "The backend shall fall back to local rule-based analysis when Ollama is unavailable."],
             ["FR-042", "The analysis result shall include safety score, verdict, health score, and health verdict."],
             ["FR-043", "The analysis result shall include flagged ingredients with severity, reason, scientific name, and personalized warning."],
             ["FR-044", "The analysis result shall include all ingredient explanations."],
             ["FR-045", "The analysis result shall include sugar, sodium, fitness, daily use, and weekly use advice."],
             ["FR-046", "The analysis result shall include healthy, gym friendly, and weight loss friendly flags."],
             ["FR-047", "The AI prompt shall require valid JSON output."],
             ["FR-048", "The backend shall normalize AI output to protect the frontend from missing fields."]])

    feature(doc, 7, "Analysis Report Display",
            "The frontend displays the complete result in a user-friendly report page.",
            "High",
            [["1", "Analysis result is returned.", "Frontend opens or updates the analysis report."],
             ["2", "User views report.", "System displays scores, verdicts, ingredients, nutrition, warnings, and recommendation."],
             ["3", "User expands ingredient card.", "System shows detailed explanation in easy language."],
             ["4", "Fallback analysis was used.", "System warns the user that the result may be less complete."]],
            [["FR-049", "The report shall show product name, brand, barcode, categories, and image when available."],
             ["FR-050", "The report shall show health score and verdict through a visual meter."],
             ["FR-051", "The report shall show ingredient cards with expandable details."],
             ["FR-052", "The ingredient card shall show why it matters for the user as one natural paragraph."],
             ["FR-053", "The report shall show nutriment breakdown for carbohydrates, protein, fat, sugar, sodium, and fiber."],
             ["FR-054", "The report shall show AI summary and recommendation."],
             ["FR-055", "The report shall warn when fallback or local rules were used."],
             ["FR-056", "The frontend shall allow saving analysis to local history."],
             ["FR-057", "The frontend shall allow sharing a report as an image."]])

    feature(doc, 8, "Scan History",
            "Authenticated users can retrieve, search, filter, and delete scan history.",
            "Medium",
            [["1", "User opens history page.", "System retrieves paginated history for the authenticated user."],
             ["2", "User searches or filters.", "System returns matching records."],
             ["3", "User deletes one scan.", "System removes that scan if it belongs to the user."],
             ["4", "User deletes all scans.", "System requires email confirmation before deletion."]],
            [["FR-058", "The backend shall return paginated scan history for the authenticated user."],
             ["FR-059", "The backend shall support filtering history by verdict."],
             ["FR-060", "The backend shall support searching history by product name."],
             ["FR-061", "The backend shall allow deletion of a single scan."],
             ["FR-062", "The backend shall allow deletion of all scans only when confirmation matches the user's email."],
             ["FR-063", "The frontend shall persist recent local history for quick access."]])

    feature(doc, 9, "System Health Monitoring",
            "The backend exposes service health information for maintainers and deployment checks.",
            "Medium",
            [["1", "Maintainer or deployment platform calls health endpoint.", "System checks service status."],
             ["2", "Dependency is unavailable.", "System reports that dependency status clearly."],
             ["3", "All checks complete.", "System returns timestamped health data."]],
            [["FR-064", "The backend shall expose /api/health/."],
             ["FR-065", "The health endpoint shall report database status."],
             ["FR-066", "The health endpoint shall report Redis status."],
             ["FR-067", "The health endpoint shall report Open Food Facts availability."],
             ["FR-068", "The health endpoint shall include a timestamp."]])

    heading(doc, 1, "5. Other Nonfunctional Requirements")
    heading(doc, 2, "5.1 Performance Requirements")
    table(doc, ["ID", "Requirement"], [
        ["NFR-001", "The frontend shall provide immediate validation feedback for barcode length and format."],
        ["NFR-002", "Product lookup results should be cached for 24 hours when Redis is available."],
        ["NFR-003", "The UI shall display progress feedback for long-running OCR and AI analysis operations."],
        ["NFR-004", "The Ollama request timeout shall be configurable."],
        ["NFR-005", "Normal page navigation should feel immediate on a modern browser and network."],
    ])

    heading(doc, 2, "5.2 Safety Requirements")
    table(doc, ["ID", "Requirement"], [
        ["NFR-006", "The system shall clearly communicate that analysis is educational and not medical advice."],
        ["NFR-007", "The system shall warn users when allergies, pregnancy, chronic health conditions, or serious dietary risks are detected."],
        ["NFR-008", "The system shall avoid presenting AI output as a guaranteed clinical decision."],
        ["NFR-009", "The system shall return safe fallback messages when product data, OCR output, or AI output is incomplete."],
    ])

    heading(doc, 2, "5.3 Security Requirements")
    table(doc, ["ID", "Requirement"], [
        ["NFR-010", "Passwords shall be hashed using bcrypt before storage."],
        ["NFR-011", "Access tokens shall be JWT-based and time-limited."],
        ["NFR-012", "Refresh tokens shall be stored in HTTP-only cookies."],
        ["NFR-013", "User scan history shall be accessible only to the owning authenticated user."],
        ["NFR-014", "Backend CORS shall allow configured frontend origins only."],
        ["NFR-015", "File upload shall reject non-image files and files larger than 10 MB."],
        ["NFR-016", "Secret keys and API keys shall be provided through environment variables in production."],
    ])

    heading(doc, 2, "5.4 Software Quality Attributes")
    table(doc, ["Attribute", "Requirement"], [
        ["Availability", "The system shall provide fallback analysis when external AI providers are unavailable."],
        ["Reliability", "The frontend shall not crash when optional AI fields are missing."],
        ["Usability", "The system shall support manual, camera, upload, and pasted-text workflows with clear messages."],
        ["Maintainability", "Backend routers shall remain separated by domain: auth, profile, scanner, analysis, history, chatbot, and health."],
        ["Portability", "The frontend shall build into static assets and the backend shall run as a Docker service."],
        ["Robustness", "The backend shall return structured HTTP errors for invalid input and unavailable product data."],
        ["Testability", "Requirement IDs shall support mapping features to test cases."],
    ])

    heading(doc, 2, "5.5 Business Rules")
    table(doc, ["Rule ID", "Business Rule"], [
        ["BR-001", "Only the owner of authenticated scan history may view or delete those records."],
        ["BR-002", "All users may analyze ingredient text, but saved server-side history requires authentication."],
        ["BR-003", "Ingredient and nutrition guidance shall remain informational and shall not claim to treat, diagnose, or prevent disease."],
        ["BR-004", "The system shall prefer product-specific data from Open Food Facts when barcode data is available."],
        ["BR-005", "External AI should be used when available, but the project must still return a usable result through fallback logic."],
    ])

    heading(doc, 1, "6. Other Requirements")
    heading(doc, 2, "6.1 Database Requirements")
    table(doc, ["Entity", "Main Fields"], [
        ["User", "id, email, hashed_password, full_name, is_active, created_at, updated_at"],
        ["UserProfile", "id, user_id, allergies, health_conditions, diet_type, updated_at"],
        ["ScanHistory", "id, user_id, barcode, product_name, product_image_url, safety_score, verdict, flagged_count, raw_product_data, created_at"],
        ["Analysis", "id, scan_id, safety_score, verdict, flagged_ingredients, all_ingredients, nutriments, ai_summary, ai_recommendation, personalized_warnings, created_at"],
    ])

    heading(doc, 2, "6.2 Data Validation Requirements")
    bullets(doc, [
        "Email must match a valid email pattern.",
        "Password must be at least 8 characters.",
        "Barcode must be 8 to 13 digits.",
        "Ingredient text must not be empty.",
        "Uploaded file must be an image and must not exceed 10 MB.",
        "Health profile values must match allowed condition, allergy, and diet lists.",
    ])

    heading(doc, 2, "6.3 Data Retention Requirements")
    bullets(doc, [
        "Authenticated scan history is stored in the backend database until deleted by the user.",
        "Local frontend scan history is stored in browser storage and limited to recent items.",
        "Uploaded OCR files are stored temporarily and should be cleaned periodically in production.",
        "Redis product cache entries are intended to expire after 24 hours.",
    ])

    heading(doc, 2, "6.4 Deployment Requirements")
    bullets(doc, [
        "Backend shall run with FastAPI and Uvicorn.",
        "Frontend shall run with Vite during development and shall build to static files for production.",
        "Docker Compose shall support local multi-service execution.",
        "Render deployment shall be supported through render.yaml.",
        "Required production environment variables include SECRET_KEY, AI_PROVIDER, GROQ_API_KEY, DATABASE_URL, REDIS_URL, ALLOWED_ORIGINS, PRODUCTION_DOMAIN, and VITE_API_BASE_URL.",
    ])

    heading(doc, 2, "6.5 Future Enhancement Requirements")
    bullets(doc, [
        "Add PDF export for analysis reports.",
        "Add admin dashboard for system monitoring.",
        "Add multilingual OCR and translation.",
        "Add serving-size based scoring in addition to per-100g scoring.",
        "Add PostgreSQL support for larger deployments.",
        "Add scheduled cleanup for uploaded OCR files.",
        "Add automated backend and frontend test suites.",
        "Add nutrition source citations in the user report.",
    ])

    heading(doc, 1, "Appendix A: Glossary")
    table(doc, ["Term", "Meaning"], [
        ["AI", "Artificial intelligence used to explain ingredients and nutrition."],
        ["API", "Application Programming Interface used for frontend-backend communication."],
        ["Barcode", "UPC or EAN numeric code used to identify a product."],
        ["CORS", "Browser security rule that controls which frontend origins can access the backend."],
        ["Groq", "External LLM provider used as primary AI analysis service."],
        ["JWT", "JSON Web Token used for access authentication."],
        ["OCR", "Optical Character Recognition used to extract text from label images."],
        ["Ollama", "Local LLM runtime used as fallback AI provider."],
        ["Open Food Facts", "Open product database used for barcode lookup."],
        ["Redis", "Cache service used for product lookup results."],
        ["SAFE", "Verdict meaning no major concern was detected in available data."],
        ["CAUTION", "Verdict meaning moderation or profile-based attention is needed."],
        ["DANGER", "Verdict meaning significant allergy, ingredient, or nutrition concern was detected."],
    ])

    heading(doc, 1, "Appendix B: Analysis Models")
    heading(doc, 2, "B.1 High-Level Data Flow")
    numbered(doc, [
        "User enters product input through barcode, image upload, pasted text, or produce selection.",
        "Frontend validates basic input and sends request to the backend.",
        "Backend fetches product data, extracts OCR text, or cleans pasted ingredient text.",
        "Backend combines ingredient text, nutrition data, and user profile.",
        "AI service analyzes the data using Groq, Ollama, or local rules.",
        "Backend stores scan and analysis results when needed.",
        "Frontend displays the report and saves recent local history.",
    ])
    heading(doc, 2, "B.2 Requirement Traceability Summary")
    table(doc, ["Feature Area", "Requirement IDs"], [
        ["Authentication", "FR-001 to FR-010"],
        ["Profile", "FR-011 to FR-016"],
        ["Barcode", "FR-017 to FR-026"],
        ["OCR Upload", "FR-027 to FR-033"],
        ["Text Analysis", "FR-034 to FR-037"],
        ["AI Analysis", "FR-038 to FR-048"],
        ["Report UI", "FR-049 to FR-057"],
        ["History", "FR-058 to FR-063"],
        ["Health Monitoring", "FR-064 to FR-068"],
        ["Nonfunctional", "NFR-001 to NFR-016 plus quality attributes and business rules"],
    ])

    heading(doc, 1, "Appendix C: To Be Determined List")
    table(doc, ["TBD ID", "Item", "Status"], [
        ["TBD-001", "Final production domain and hosting URLs.", "To be confirmed during deployment."],
        ["TBD-002", "Final AI model choice if Groq model changes after testing.", "To be confirmed during final evaluation."],
        ["TBD-003", "Production cleanup schedule for uploaded OCR files.", "To be defined before long-term hosting."],
        ["TBD-004", "Whether PostgreSQL will replace SQLite for high-traffic production use.", "Future enhancement."],
    ])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
