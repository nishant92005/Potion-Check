from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path("docs/PotionCheck_ProjectFlow.docx")
PROJECT = "PotionCheck"
TEAM = "Nishant (2310993892)  |  Mohit (2310993967)"
PROGRAM = "BE-CSE (AI)  |  Semester 6  |  Chitkara University"


def set_margins(section):
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)


def page_field(paragraph):
    run = paragraph.add_run("Page ")
    run.bold = True
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def setup_header_footer(section):
    set_margins(section)
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.text = ""
    t = header.add_table(rows=1, cols=2, width=Inches(6.83))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(3.5)
    t.columns[1].width = Inches(3.33)
    left = t.cell(0, 0).paragraphs[0]
    right = t.cell(0, 1).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = left.add_run(f"{PROJECT}  —  Project Flow Document")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r = right.add_run(PROGRAM)
    r.font.name = "Arial"
    r.font.size = Pt(8)

    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.text = ""
    ft = footer.add_table(rows=1, cols=2, width=Inches(6.83))
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    ft.autofit = False
    ft.columns[0].width = Inches(5.8)
    ft.columns[1].width = Inches(1.03)
    left = ft.cell(0, 0).paragraphs[0]
    right = ft.cell(0, 1).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = left.add_run(TEAM)
    r.font.name = "Arial"
    r.font.size = Pt(7.5)
    page_field(right)
    for run in right.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)


def style_doc(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for name, size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 10.5)]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(7)


def para(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=9.5, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def heading(doc, level, text):
    doc.add_heading(text, level=level)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.name = "Arial"
        r.font.size = Pt(9.5)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.name = "Arial"
        r.font.size = Pt(9.5)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    if widths:
        for i, width in enumerate(widths):
            t.columns[i].width = Inches(width)
    for i, h in enumerate(headers):
        shade(t.rows[0].cells[i], "EAF2F8")
        cell_text(t.rows[0].cells[i], h, bold=True, size=9)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)
    doc.add_paragraph()
    return t


def phase_block(doc, number, name, period, owner, deliverable, steps):
    table(doc, ["Phase", str(number), name, period], [["Owner:", owner, "Key Deliverable:", deliverable]], [0.7, 0.7, 3.1, 2.0])
    numbered(doc, steps)


def workflow_table(doc, title, steps):
    heading(doc, 2, title)
    rows = []
    for idx, (label, desc) in enumerate(steps, 1):
        rows.append([idx, label, desc])
        if idx != len(steps):
            rows.append(["↓", "", ""])
    table(doc, ["#", "Process", "Description"], rows, [0.5, 2.0, 4.3])


def build():
    doc = Document()
    style_doc(doc)
    setup_header_footer(doc.sections[0])

    doc.add_paragraph()
    doc.add_paragraph()
    para(doc, PROJECT, WD_ALIGN_PARAGRAPH.CENTER, 26, True)
    para(doc, "Project Flow Document", WD_ALIGN_PARAGRAPH.CENTER, 17, True)
    para(doc, "AI Ingredient Intelligence Scanner & Personalized Food Analysis Platform", WD_ALIGN_PARAGRAPH.CENTER, 11, italic=True)
    doc.add_paragraph()

    cover_rows = [
        ["Project Name", PROJECT],
        ["Version", "1.0"],
        ["Document Type", "Project Flow Document"],
        ["Prepared By", "Nishant (Roll No. 2310993892)\nMohit (Roll No. 2310993967)"],
        ["Institution", "Chitkara University Institute of Engineering & Technology, Rajpura"],
        ["Program", "BE-CSE (Artificial Intelligence) — Semester 6"],
        ["Date", "May 2026"],
        ["SRS Reference", "PotionCheck SRS v1.0 (May 2026)"],
        ["Status", "Final Submission"],
    ]
    t = table(doc, ["", ""], cover_rows, [1.8, 4.8])
    t.rows[0]._element.getparent().remove(t.rows[0]._element)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    setup_header_footer(doc.sections[-1])
    heading(doc, 1, "1. Introduction")
    para(doc, "PotionCheck is a web-based, AI-powered ingredient intelligence scanner developed as a B.Tech Semester 6 project at Chitkara University, Rajpura, Punjab. This Project Flow Document describes the complete lifecycle of the project from planning and requirements to development, testing, documentation, and final delivery. It explains how the system is built, how food data flows through the application, and how each phase connects with the requirements defined in PotionCheck SRS v1.0.")
    heading(doc, 2, "1.1 Purpose")
    para(doc, "This document provides:")
    bullets(doc, [
        "A clear overview of all project phases with timelines and milestones.",
        "The complete data and process flow across frontend, backend, AI, OCR, database, and external services.",
        "Traceability from functional requirements to implementation phases.",
        "AI analysis workflow and system architecture diagrams in document form.",
        "Roles and responsibilities of all team members.",
    ])
    heading(doc, 2, "1.2 System Overview")
    para(doc, "PotionCheck solves the problem of confusing packaged food labels. Many users do not understand additives, nutrition values, allergens, or ingredient risks written on product packaging. PotionCheck allows users to scan a barcode, upload a food label image, paste ingredient text, or select a product flow and then receive a simple report. The report includes safety score, health score, verdict, flagged ingredients, nutrition observations, personalized warnings, and recommendation.")
    para(doc, "The platform is built with a React and Vite frontend, a FastAPI backend, SQLite database storage, Redis caching, Open Food Facts barcode lookup, OCR support, and AI analysis. Groq is used as the primary LLM provider. Ollama is used as fallback, and local rule-based analysis is available when external AI services are not reachable. Authentication is handled with JWT access tokens, HTTP-only refresh cookies, and bcrypt password hashing.")

    heading(doc, 1, "2. Project Phases & Timeline")
    para(doc, "The project is structured into seven sequential phases spanning January 2026 to May 2026. Each phase has defined entry criteria, deliverables, and exit criteria.")
    table(doc, ["#", "Phase", "Period", "Owner", "Key Deliverable"], [
        ["1", "Planning & Requirements", "Jan 2026", "All Members", "SRS v1.0, Project Charter"],
        ["2", "System Design & Architecture", "Feb 2026", "All Members", "Architecture Diagram, DB Schema, API Design"],
        ["3", "Backend Development", "Feb – Mar 2026", "Nishant", "FastAPI API, Auth, DB, AI/OCR Integration"],
        ["4", "Frontend Development", "Mar – Apr 2026", "Mohit", "React UI, Scanner, Analysis, Profile, History"],
        ["5", "AI, OCR & Product Lookup Integration", "Apr 2026", "Nishant + Mohit", "Barcode Lookup, OCR, Groq/Ollama Analysis"],
        ["6", "Testing & QA", "Apr – May 2026", "All Members", "Test Cases, Bug Fixes, Security Checks"],
        ["7", "Documentation & Delivery", "May 2026", "All Members", "Project Report, Flow Doc, Presentation, Submission"],
    ], [0.35, 1.55, 1.0, 1.1, 2.7])

    heading(doc, 1, "3. Detailed Phase Flow")
    phase_block(doc, 1, "Planning & Requirements", "January 2026", "All Members", "SRS v1.0, Project Charter", [
        "Conduct kickoff meetings and assign responsibilities: Nishant for backend, database, AI integration, OCR, and deployment; Mohit for frontend, UI flow, testing support, and documentation.",
        "Identify target users such as health-conscious consumers, people with allergies, gym users, students, and general users who want to understand food labels.",
        "Define the core problem: packaged food labels are hard to understand and users need simple personalized explanations.",
        "Write and review SRS v1.0 with functional requirements, non-functional requirements, assumptions, constraints, and limitations.",
        "Set up GitHub repository, development environment, Python backend, Node.js frontend, Docker files, and environment variable structure.",
        "Finalize project scope, milestones, team responsibilities, and risk register.",
    ])
    phase_block(doc, 2, "System Design & Architecture", "February 2026", "All Members", "Architecture Diagram, DB Schema, API Design", [
        "Design three-layer architecture: React/Vite frontend → FastAPI backend → SQLite/Redis/Data and external services.",
        "Design product analysis pipeline: User Input → Scanner/OCR/Text Parser → Product Data → User Profile → AI Analysis → Report UI → History.",
        "Design REST API contract including endpoints, request/response schemas, authentication flow, and error codes.",
        "Design database schema: Users, UserProfile, ScanHistory, and Analysis tables.",
        "Design fallback architecture: Groq primary AI → Ollama fallback → local rule-based analysis.",
        "Document technology decisions: React, Vite, Tailwind CSS, FastAPI, SQLite, Redis, Open Food Facts, Tesseract OCR, Groq, Ollama, Docker, and Render.",
    ])
    phase_block(doc, 3, "Backend Development", "February – March 2026", "Nishant", "FastAPI API, JWT Auth, Database, AI and OCR Integration", [
        "Initialize FastAPI backend project and configure CORS, environment variables, database initialization, and router structure.",
        "Implement JWT-based authentication: user registration, login, refresh token, logout, and current user endpoint.",
        "Hash all passwords with bcrypt before storage and validate email format and password length during registration.",
        "Implement SQLite database connection through SQLAlchemy async models for User, UserProfile, ScanHistory, and Analysis.",
        "Build scanner endpoints for barcode lookup, barcode analysis, image upload, OCR extraction, and pasted text parsing.",
        "Integrate Open Food Facts API for product lookup by barcode and add Redis caching for repeated product requests.",
        "Implement AI analysis service using Groq as primary provider, Ollama fallback, and local rule-based fallback.",
        "Implement scan history APIs with pagination, verdict filtering, product-name search, single deletion, and all-scan deletion with email confirmation.",
        "Implement health check endpoint to report backend, database, Redis, and Open Food Facts status.",
    ])
    phase_block(doc, 4, "Frontend Development", "March – April 2026", "Mohit", "React UI, Scanner, Analysis Report, Profile, History", [
        "Initialize React + Vite + Tailwind CSS frontend project and configure environment variables.",
        "Build Landing page with application overview, visual background, navigation, and clear entry points.",
        "Build Profile page for registration, login, saved health profile, allergies, health conditions, and diet type.",
        "Build Scanner page with barcode scanning, manual barcode entry, uploaded label image, pasted ingredient text, and scan status feedback.",
        "Build Analysis page with product details, score meter, verdict, nutriment breakdown, flagged ingredients, all ingredient explanations, AI summary, and recommendation.",
        "Build History page for saved scan records, local recent history, filters, searching, and deletion actions.",
        "Build Chatbot page for ingredient and food-related questions backed by project AI/RAG service.",
        "Implement responsive layout, animated background, navbar, toast notifications, page transitions, and mobile-friendly interactions.",
        "Centralize API communication in frontend service modules for scanner, profile, auth, analysis, history, and chatbot.",
    ])
    phase_block(doc, 5, "AI, OCR & Product Lookup Integration", "April 2026", "Nishant + Mohit", "Barcode Lookup, OCR, Groq/Ollama Analysis, Report Integration", [
        "Integrate barcode flow end-to-end: frontend scanner → backend barcode endpoint → Open Food Facts → cache → analysis report.",
        "Validate product data handling when ingredients, nutriments, brand, category, or image fields are missing.",
        "Integrate label upload flow end-to-end: image upload → backend validation → OCR extraction → editable text → AI analysis.",
        "Integrate pasted text flow end-to-end: user text → parser → ingredient list preview → custom analysis.",
        "Validate AI fallback flow: Groq response, Groq failure, Ollama fallback, and local rule-based analysis.",
        "Normalize AI JSON output so frontend receives stable fields for safety score, verdict, health score, warnings, ingredients, summary, and recommendation.",
        "Perform cross-browser testing on Chrome, Edge, Firefox, and mobile viewports.",
    ])
    phase_block(doc, 6, "Testing & QA", "April – May 2026", "All Members", "Test Cases, Bug Fixes, Security Checks", [
        "Execute test cases across Authentication, Profile, Scanner, OCR Upload, Text Analysis, AI Analysis, Report UI, History, Chatbot, Health Check, and Deployment modules.",
        "Test invalid inputs including duplicate email, wrong password, invalid barcode, empty ingredient text, non-image upload, and oversized image upload.",
        "Validate JWT token issuance, refresh handling, logout, protected endpoints, and rejection of unauthorized access.",
        "Conduct CORS and file-upload checks and verify sensitive keys are loaded through environment variables.",
        "Measure user flow performance and confirm progress indicators appear during lookup, OCR, and AI analysis.",
        "Log defects, triage by severity, and fix critical and high-severity issues before final submission.",
        "Review UI on desktop and mobile screens to ensure content is clear, readable, and not overlapping.",
    ])
    phase_block(doc, 7, "Documentation & Delivery", "May 2026", "All Members", "Project Report, SRS, Flow Document, Test Cases, Presentation, Final Submission", [
        "Finalize SRS document using IEEE-style structure with all functional and non-functional requirements.",
        "Complete Project Charter with scope, milestones, resources, risks, and sign-off section.",
        "Write Project Flow Document covering all phases, data flows, AI/OCR workflow, architecture, traceability, roles, risks, stack, system requirements, and future scope.",
        "Compile test case document with module-wise test cases and status.",
        "Prepare final project presentation and live demo for academic evaluation committee.",
        "Submit all project artifacts: SRS, Charter, Flow Document, Test Cases, Project Report, Source Code, and Presentation.",
    ])

    heading(doc, 1, "4. System Architecture & Data Flow")
    heading(doc, 2, "4.1 Three-Layer Architecture")
    table(doc, ["Layer", "Technology", "Description"], [
        ["Layer 1 — Frontend (SPA)", "React + Vite + Tailwind CSS", "Responsive single-page application. Handles user interaction, scanner UI, profile, analysis report, history, chatbot, and API calls."],
        ["Layer 2 — Backend (API)", "FastAPI + Uvicorn", "Handles authentication, profile, scanner requests, OCR upload, AI analysis, history, chatbot, health checks, validation, and errors."],
        ["Layer 3 — Data & External Services", "SQLite + Redis + Open Food Facts + Groq + Ollama + OCR", "Stores users and analyses, caches product data, fetches barcode products, extracts label text, and generates AI reports."],
    ], [1.6, 1.8, 3.4])
    heading(doc, 2, "4.2 Request Flow — Barcode Product Analysis")
    numbered(doc, [
        "User scans a barcode using camera, uploads a barcode image, or enters barcode manually on the Scanner page.",
        "Frontend validates the barcode length and format before sending the request.",
        "Frontend sends POST /api/scanner/barcode/analyze with barcode and optional profile data to the FastAPI backend.",
        "Backend checks Redis cache for the product. If cache is empty, backend fetches product data from Open Food Facts.",
        "Backend returns 404 if product is not found and 422 if product exists but ingredient data is missing.",
        "Backend combines product ingredients, nutrition values, product context, and user profile.",
        "AI service analyzes the product using Groq. If Groq is unavailable, it uses Ollama. If Ollama is unavailable, it uses local rules.",
        "Backend saves scan history and analysis result in SQLite for authenticated users.",
        "Backend returns JSON response containing product details, scores, verdict, flagged ingredients, nutriments, summary, and recommendation.",
        "Frontend renders the Analysis page with visual score, ingredient cards, warnings, nutrition breakdown, and user-friendly explanation.",
    ])
    heading(doc, 2, "4.3 Authentication Flow")
    numbered(doc, [
        "User fills Sign Up form with full name, email, and password. Frontend validates basic form input.",
        "Frontend sends POST /api/auth/register to the backend.",
        "Backend validates email uniqueness and password length.",
        "Backend hashes password using bcrypt and stores the new user in SQLite.",
        "On login, frontend sends POST /api/auth/login with user credentials.",
        "Backend verifies credentials against the hashed password.",
        "Backend issues a signed JWT access token and an HTTP-only refresh cookie.",
        "Frontend sends the access token for protected API requests.",
        "Backend validates JWT on protected endpoints and rejects expired or invalid tokens with HTTP 401.",
        "User can refresh token with POST /api/auth/refresh and log out with POST /api/auth/logout.",
    ])
    heading(doc, 2, "4.4 OCR and Custom Text Analysis Flow")
    numbered(doc, [
        "User uploads a food label image or pastes ingredient text on the Scanner page.",
        "For image upload, backend validates that file type is image and file size is not more than 10 MB.",
        "Backend stores the image temporarily and runs OCR to extract ingredient text.",
        "Frontend displays extracted text so the user can review and correct it if needed.",
        "For pasted text, backend cleans and parses the ingredient list.",
        "User sends the final ingredient text for analysis.",
        "Backend combines text with user profile and available nutrition data.",
        "AI service returns safety score, health score, verdict, ingredient explanations, warnings, summary, and recommendation.",
        "Frontend displays the final report and saves recent local history.",
    ])

    heading(doc, 1, "5. AI Analysis Workflow")
    para(doc, "PotionCheck operates through a structured ingredient analysis workflow. The workflow combines user input, product data, profile personalization, external knowledge, AI reasoning, fallback safety, and report rendering.")
    workflow_table(doc, "5.1 Ingredient Intelligence Pipeline", [
        ("User Input", "User scans barcode, uploads label image, pastes ingredients, or opens a previous analysis."),
        ("Input Processor", "Frontend and backend validate barcode, image, or text input and show clear status messages."),
        ("Product Data Layer", "Backend fetches product data from Open Food Facts or extracts text using OCR."),
        ("Profile Personalizer", "Backend applies allergies, health conditions, and diet type from saved or supplied user profile."),
        ("AI Provider Chain", "Groq performs primary analysis; Ollama and local rules provide fallback analysis."),
        ("Result Normalizer", "Backend normalizes AI output into stable JSON fields required by the frontend."),
        ("Storage Layer", "ScanHistory and Analysis records are saved in SQLite when applicable."),
        ("Response to UI", "Frontend renders score, verdict, flagged ingredients, nutriments, warnings, and recommendations."),
    ])
    para(doc, "Left flow: User Input → Scanner/OCR/Text Processing → Product Data → AI Analysis. Right flow: Profile Personalization → Result Normalization → Storage → Report UI.")

    heading(doc, 1, "6. Functional Requirements Traceability")
    para(doc, "All functional requirements from PotionCheck SRS v1.0 are mapped to their implementation phases below.")
    table(doc, ["Req ID", "Module", "Description", "Phase", "Priority"], [
        ["REQ-01", "Authentication", "User registration with unique email, full name, and password.", "Phase 3", "Must"],
        ["REQ-02", "Authentication", "Password length and email format validation.", "Phase 3", "Must"],
        ["REQ-03", "Authentication", "Password hashing with bcrypt before database storage.", "Phase 3", "Must"],
        ["REQ-04", "Authentication", "JWT access token and HTTP-only refresh cookie issuance.", "Phase 3", "Must"],
        ["REQ-05", "Authentication", "Protected current-user and user-specific endpoints.", "Phase 3", "Must"],
        ["REQ-06", "Profile", "Save allergies, health conditions, and diet type.", "Phase 3 & 4", "Must"],
        ["REQ-07", "Barcode Scanner", "Manual barcode entry and live camera scanning.", "Phase 4 & 5", "Must"],
        ["REQ-08", "Barcode Scanner", "Uploaded barcode image support.", "Phase 4 & 5", "Should"],
        ["REQ-09", "Product Lookup", "Fetch product data from Open Food Facts.", "Phase 3 & 5", "Must"],
        ["REQ-10", "Product Lookup", "Cache product lookup result using Redis.", "Phase 3 & 5", "Should"],
        ["REQ-11", "OCR Upload", "Accept image upload and reject invalid or oversized files.", "Phase 3 & 5", "Must"],
        ["REQ-12", "OCR Upload", "Extract ingredient text from uploaded label image.", "Phase 3 & 5", "Must"],
        ["REQ-13", "Text Analysis", "Clean and parse pasted ingredient text.", "Phase 3 & 4", "Must"],
        ["REQ-14", "AI Analysis", "Analyze ingredients using Groq primary provider.", "Phase 3 & 5", "Must"],
        ["REQ-15", "AI Analysis", "Fallback to Ollama and local rule-based analysis.", "Phase 3 & 5", "Must"],
        ["REQ-16", "AI Analysis", "Return safety score, health score, verdict, flagged ingredients, and recommendation.", "Phase 3 & 5", "Must"],
        ["REQ-17", "Report UI", "Display product details, nutrition, warnings, explanations, and score meter.", "Phase 4 & 5", "Must"],
        ["REQ-18", "Report UI", "Show easy-language explanation for why each concern matters.", "Phase 4 & 5", "Must"],
        ["REQ-19", "History", "Store and retrieve authenticated user's scan history.", "Phase 3 & 4", "Must"],
        ["REQ-20", "History", "Search, filter, delete one scan, and delete all scans with confirmation.", "Phase 3 & 4", "Should"],
        ["REQ-21", "Chatbot", "Provide ingredient and food question support through chatbot service.", "Phase 3 & 4", "Should"],
        ["REQ-22", "Health Check", "Expose backend health endpoint with dependency status.", "Phase 3", "Must"],
        ["REQ-23", "Deployment", "Support Docker Compose and Render deployment.", "Phase 7", "Must"],
        ["REQ-24", "Documentation", "Provide SRS, Flow Document, project report, presentation, and source code.", "Phase 7", "Must"],
    ], [0.65, 1.15, 2.8, 1.05, 0.8])

    heading(doc, 1, "7. Roles & Responsibilities")
    table(doc, ["Team Member", "Roll No.", "Role", "Responsibilities"], [
        ["Nishant", "2310993892", "Backend / AI Integration Lead", "FastAPI backend, authentication, database models, scanner APIs, Open Food Facts integration, Groq/Ollama AI analysis, OCR, health checks, deployment configuration."],
        ["Mohit", "2310993967", "Frontend Developer / UI-UX / QA Support", "React/Vite frontend, scanner page, analysis report UI, profile and history views, chatbot page, responsive design, API service modules, UI testing, documentation support."],
    ], [1.0, 0.9, 1.55, 3.25])
    heading(doc, 2, "7.1 Communication Protocol")
    table(doc, ["Channel", "Purpose", "Frequency"], [
        ["GitHub", "Source code version control, commits, issue tracking, and code review.", "Daily"],
        ["WhatsApp / Discord", "Team coordination, quick updates, and blocker resolution.", "Daily"],
        ["Weekly Review", "Progress check, milestone validation, and task reassignment.", "Weekly"],
        ["Academic Supervisor", "Faculty sign-off on deliverables and evaluation checkpoints.", "As required"],
    ], [1.4, 4.2, 1.1])

    heading(doc, 1, "8. Risk Register & Mitigation")
    table(doc, ["ID", "Risk", "Impact", "Likelihood", "Mitigation Strategy"], [
        ["R-01", "Groq API downtime or quota limit disrupts AI analysis.", "High", "Medium", "Use Ollama fallback and local rule-based analysis for graceful degradation."],
        ["R-02", "Open Food Facts product data is missing or inaccurate.", "Medium", "High", "Show clear error or missing-data message and allow pasted text or label upload as alternate input."],
        ["R-03", "OCR accuracy is poor due to blurry label image.", "Medium", "Medium", "Allow user to review and edit extracted text before final analysis."],
        ["R-04", "Camera permission or browser limitation blocks live barcode scan.", "Medium", "Medium", "Provide manual barcode entry and image upload as alternate workflows."],
        ["R-05", "Users treat the system as medical advice.", "High", "Low", "Display educational-use disclaimer and advise professional consultation for serious health concerns."],
        ["R-06", "JWT token or account security issue.", "High", "Low", "Use hashed passwords, short-lived access tokens, HTTP-only refresh cookies, and protected endpoints."],
        ["R-07", "SQLite becomes insufficient for high-traffic production.", "Medium", "Low", "Keep schema portable and plan PostgreSQL migration for future production scaling."],
        ["R-08", "Team member unavailability due to academic workload.", "Medium", "Medium", "Share documentation, keep code organized, and review progress weekly."],
    ], [0.55, 2.1, 0.8, 0.9, 2.5])

    heading(doc, 1, "9. Technology Stack Summary")
    table(doc, ["Layer / Component", "Technology", "Purpose"], [
        ["Frontend Framework", "React 18 + Vite", "Single-page application, routing, and component architecture."],
        ["Frontend Styling", "Tailwind CSS", "Responsive styling and modern UI layout."],
        ["Frontend State", "Zustand", "Profile, scanner, history, and app state management."],
        ["HTTP Client", "Axios", "Frontend API request handling."],
        ["Barcode", "html5-qrcode", "Live camera and image-based barcode scanning."],
        ["OCR", "Tesseract.js / pytesseract", "Ingredient text extraction from label images."],
        ["Backend Runtime", "Python 3 + Uvicorn", "Backend application runtime server."],
        ["Backend Framework", "FastAPI", "REST API routing, validation, dependency injection, and middleware."],
        ["Database", "SQLite + SQLAlchemy async", "Stores users, profiles, scan history, and analyses."],
        ["Cache", "Redis", "Caches product lookup data for faster repeated requests."],
        ["AI API", "Groq", "Primary LLM ingredient analysis provider."],
        ["AI Fallback", "Ollama + local rules", "Fallback analysis when Groq is unavailable."],
        ["External Product API", "Open Food Facts", "Barcode product lookup and nutrition data."],
        ["Authentication", "JWT + HTTP-only refresh cookie", "Session and protected API access."],
        ["Password Security", "bcrypt", "Password hashing before database storage."],
        ["Version Control", "Git + GitHub", "Source control and collaborative development."],
        ["Dev Tools", "VS Code, Postman", "IDE and API testing/debugging."],
        ["Deployment Target", "Docker Compose + Render", "Local and cloud deployment support."],
    ], [1.7, 1.8, 3.2])
    heading(doc, 2, "9.1 System Requirements Summary")
    table(doc, ["Requirement Type", "Minimum Specification", "Recommended"], [
        ["Development CPU", "Intel Core i3 / AMD equivalent, 2.0 GHz", "Intel Core i5+ / 2.5 GHz+"],
        ["Development RAM", "4 GB", "8 GB"],
        ["Storage (Dev)", "10 GB free, SSD preferred", "20 GB SSD"],
        ["Production CPU", "2 vCPU cores", "4 vCPU cores"],
        ["Production RAM", "4 GB", "8 GB"],
        ["Production Storage", "20 GB SSD", "40 GB SSD"],
        ["Network (Client)", "5 Mbps", "10 Mbps+"],
        ["Browser (Client)", "Chrome, Firefox, Edge, Safari latest", "Latest versions"],
        ["Mobile Client", "Android 8+ / iOS 14+", "Android 12+ / iOS 16+"],
    ], [1.8, 2.55, 2.35])

    heading(doc, 1, "10. Future Scope & Roadmap")
    para(doc, "PotionCheck has strong future scope beyond the current academic project. The following enhancements are planned for future iterations:")
    table(doc, ["ID", "Feature", "Priority", "Description"], [
        ["FR-01", "PDF Export", "High", "Allow users to download complete analysis reports as PDF files."],
        ["FR-02", "Mobile Application", "High", "Native Android and iOS apps with camera-first barcode scanning and saved reports."],
        ["FR-03", "PostgreSQL Migration", "Medium", "Move from SQLite to PostgreSQL for larger production deployments."],
        ["FR-04", "Serving Size Analysis", "High", "Add serving-size based scoring in addition to per-100g nutrition scoring."],
        ["FR-05", "Multilingual OCR", "Medium", "Support Hindi and other regional languages for ingredient labels."],
        ["FR-06", "Nutrition Citations", "Medium", "Add source references and citations for health and nutrition explanations."],
        ["FR-07", "Admin Dashboard", "Medium", "Monitor system health, API errors, upload cleanup, and usage statistics."],
        ["FR-08", "Allergen Alert Mode", "High", "Provide stronger visual and audio alerts for serious allergy matches."],
        ["FR-09", "Product Comparison", "Medium", "Compare two products and recommend the healthier choice for the user's profile."],
    ], [0.65, 1.45, 0.8, 3.9])
    para(doc, f"{PROJECT}  |  Project Flow Document v1.0  |  Nishant • Mohit  |  Chitkara University  |  May 2026", WD_ALIGN_PARAGRAPH.CENTER, 8)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
